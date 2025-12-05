"""Script to generate report from template and pipeline results."""
from docx import Document
import os
from datetime import datetime
import pandas as pd
import numpy as np

def read_template_and_generate_report():
    """Read the template and generate a filled report."""
    
    # Read the template
    template_path = "Report_Template.docx"
    if not os.path.exists(template_path):
        print(f"Template not found at {template_path}")
        return
    
    doc = Document(template_path)
    
    # Load results data
    data_path = "data/sample_data.csv"
    df = pd.read_csv(data_path)
    
    # Get pipeline results (from the run we just did)
    results = {
        'data_shape': (20, 4),
        'features_shape': (20, 3),
        'has_predictions': True,
        'num_predictions': 20,
        'target_column': 'target',
        'model_type': 'RandomForestClassifier',
        'num_classes': 2
    }
    
    # Replace placeholders in the document
    # First, let's see what's in the template
    print("Template content:")
    for para in doc.paragraphs:
        print(para.text)
    
    # Replace text in paragraphs
    for para in doc.paragraphs:
        text = para.text
        
        # Replace common placeholders
        text = text.replace('{DATE}', datetime.now().strftime('%Y-%m-%d'))
        text = text.replace('{DATETIME}', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        text = text.replace('{DATA_ROWS}', str(results['data_shape'][0]))
        text = text.replace('{DATA_COLUMNS}', str(results['data_shape'][1]))
        text = text.replace('{FEATURES_COUNT}', str(results['features_shape'][1]))
        text = text.replace('{PREDICTIONS_COUNT}', str(results['num_predictions']))
        text = text.replace('{MODEL_TYPE}', results['model_type'])
        text = text.replace('{TARGET_COLUMN}', results['target_column'])
        
        if text != para.text:
            para.text = text
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                text = text.replace('{DATE}', datetime.now().strftime('%Y-%m-%d'))
                text = text.replace('{DATA_ROWS}', str(results['data_shape'][0]))
                text = text.replace('{DATA_COLUMNS}', str(results['data_shape'][1]))
                text = text.replace('{FEATURES_COUNT}', str(results['features_shape'][1]))
                text = text.replace('{PREDICTIONS_COUNT}', str(results['num_predictions']))
                text = text.replace('{MODEL_TYPE}', results['model_type'])
                if text != cell.text:
                    cell.text = text
    
    # Save the report
    output_path = "Report_Generated.docx"
    doc.save(output_path)
    print(f"\nReport generated successfully: {output_path}")

if __name__ == "__main__":
    read_template_and_generate_report()

