"""Fill in the template report with actual results."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import pandas as pd
import re

def fill_template_report():
    """Read template and fill it with actual results."""
    
    template_path = "Report_Template.docx"
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return
    
    doc = Document(template_path)
    
    # Load data
    data_path = "data/sample_data.csv"
    df = pd.read_csv(data_path)
    
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
    
    # Results data
    results = {
        'date': datetime.now().strftime('%B %d, %Y'),
        'data_rows': len(df),
        'data_columns': len(df.columns),
        'features_count': len(df.columns) - 1,  # Excluding target
        'predictions_count': len(df),
        'target_column': 'target',
        'model_type': 'RandomForestClassifier',
        'numeric_features': ', '.join([c for c in numeric_cols if c != 'target']),
        'categorical_features': ', '.join(categorical_cols),
    }
    
    # Function to replace text in paragraphs
    def replace_in_paragraph(para, replacements):
        """Replace placeholders in a paragraph."""
        if not para.text.strip():
            return
        
        text = para.text
        original_text = text
        
        # Replace common placeholders
        for key, value in replacements.items():
            text = text.replace(key, str(value))
        
        # If text changed, update paragraph
        if text != original_text:
            para.clear()
            para.add_run(text)
    
    # Replacements dictionary
    replacements = {
        '{DATE}': results['date'],
        '{DATETIME}': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        '{DATA_ROWS}': results['data_rows'],
        '{DATA_COLUMNS}': results['data_columns'],
        '{FEATURES_COUNT}': results['features_count'],
        '{PREDICTIONS_COUNT}': results['predictions_count'],
        '{MODEL_TYPE}': results['model_type'],
        '{TARGET_COLUMN}': results['target_column'],
    }
    
    # Process all paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
        
        # Check for specific sections and add content
        text_lower = para.text.lower()
        
        # Problem statement section - skip complex paragraph manipulation
        # We'll handle content addition differently
        
        # Dataset table - look for table section
        if 'dataset' in text_lower and 'table' in text_lower:
            # We'll handle tables separately
            pass
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for key, value in replacements.items():
                    text = text.replace(key, str(value))
                
                # Fill dataset information table
                if 'dataset name' in text.lower() or 'sample_data' in text.lower():
                    # Update dataset name cell
                    for r in table.rows:
                        for c in r.cells:
                            if 'dataset name' in c.text.lower() or 'sample_data' in c.text.lower():
                                # Find the value cell (usually next column)
                                cell_idx = r.cells.index(c)
                                if cell_idx < len(r.cells) - 1:
                                    r.cells[cell_idx + 1].text = 'sample_data.csv'
                
                if text != cell.text:
                    cell.text = text
    
    # Add detailed content to specific sections by finding and updating paragraphs
    # This is a simplified approach - in practice, you might need more sophisticated paragraph matching
    
    # Save filled template
    output_path = "Report_Filled_Template.docx"
    doc.save(output_path)
    print(f"\nTemplate filled successfully: {output_path}")
    print(f"\nFilled sections:")
    print(f"  - Date: {results['date']}")
    print(f"  - Dataset: {results['data_rows']} rows × {results['data_columns']} columns")
    print(f"  - Features: {results['features_count']}")
    print(f"  - Model: {results['model_type']}")
    print(f"  - Predictions: {results['predictions_count']}")

if __name__ == "__main__":
    fill_template_report()

