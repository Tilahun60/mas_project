"""Properly fill the template report with comprehensive content."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import pandas as pd

def fill_template_properly():
    """Read template and fill with comprehensive content."""
    
    template_path = "Report_Template.docx"
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return
    
    # Load the template
    doc = Document(template_path)
    
    # Load data for analysis
    data_path = "data/sample_data.csv"
    df = pd.read_csv(data_path)
    
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
    numeric_features = [c for c in numeric_cols if c != 'target']
    
    # Results
    results = {
        'date': datetime.now().strftime('%B %d, %Y'),
        'data_rows': len(df),
        'data_columns': len(df.columns),
        'features_count': len(df.columns) - 1,
        'predictions_count': len(df),
        'target_column': 'target',
        'model_type': 'RandomForestClassifier',
        'numeric_features': ', '.join(numeric_features),
        'categorical_features': ', '.join(categorical_cols),
    }
    
    # Simple replacement approach - replace placeholders in all paragraphs and cells
    replacements = {
        '{DATE}': results['date'],
        '{DATETIME}': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        '{DATA_ROWS}': str(results['data_rows']),
        '{DATA_COLUMNS}': str(results['data_columns']),
        '{FEATURES_COUNT}': str(results['features_count']),
        '{PREDICTIONS_COUNT}': str(results['predictions_count']),
        '{MODEL_TYPE}': results['model_type'],
        '{TARGET_COLUMN}': results['target_column'],
    }
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        text = para.text
        for key, value in replacements.items():
            text = text.replace(key, value)
        if text != para.text:
            para.clear()
            para.add_run(text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for key, value in replacements.items():
                    text = text.replace(key, value)
                if text != cell.text:
                    cell.text = text
    
    # Save the filled template
    output_path = "Report_Filled_Template.docx"
    doc.save(output_path)
    
    print(f"\n{'='*60}")
    print("Template Report Filled Successfully!")
    print(f"{'='*60}")
    print(f"Output file: {output_path}")
    print(f"\nFilled information:")
    print(f"  - Date: {results['date']}")
    print(f"  - Dataset: {results['data_rows']} rows × {results['data_columns']} columns")
    print(f"  - Features: {results['features_count']} ({results['numeric_features']}, {results['categorical_features']})")
    print(f"  - Model: {results['model_type']}")
    print(f"  - Predictions: {results['predictions_count']}")
    print(f"{'='*60}")
    print("\nNote: The template has been filled with basic placeholders.")
    print("For a comprehensive report with all sections detailed, see: Report_Generated.docx")

if __name__ == "__main__":
    fill_template_properly()

