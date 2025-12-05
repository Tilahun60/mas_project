"""Create final comprehensive report using the template structure."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import pandas as pd

def create_final_report():
    """Create final report based on template structure with all sections filled."""
    
    # Load template
    template_path = "Report_Template.docx"
    doc = Document(template_path)
    
    # Load data
    data_path = "data/sample_data.csv"
    df = pd.read_csv(data_path)
    
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
    numeric_features = [c for c in numeric_cols if c != 'target']
    
    results = {
        'date': datetime.now().strftime('%B %d, %Y'),
        'data_rows': len(df),
        'data_columns': len(df.columns),
        'features_count': len(df.columns) - 1,
        'predictions_count': len(df),
        'target_column': 'target',
        'model_type': 'RandomForestClassifier',
        'numeric_features': ', '.join(numeric_features),
        'categorical_features': ', '.join(categorical_cols),
    }
    
    # Replace placeholders
    replacements = {
        '{DATE}': results['date'],
        '{DATETIME}': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        '{DATA_ROWS}': str(results['data_rows']),
        '{DATA_COLUMNS}': str(results['data_columns']),
        '{FEATURES_COUNT}': str(results['features_count']),
        '{PREDICTIONS_COUNT}': str(results['predictions_count']),
        '{MODEL_TYPE}': results['model_type'],
        '{TARGET_COLUMN}': results['target_column'],
    }
    
    # Process paragraphs and add detailed content
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Replace placeholders
        for key, value in replacements.items():
            text = text.replace(key, value)
        
        # Add content to specific sections
        text_lower = text.lower()
        
        # Problem statement
        if 'problem' in text_lower and 'statement' in text_lower and len(text) < 100:
            # Find the next paragraph that needs content
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if len(next_para.text.strip()) < 200:  # Likely a placeholder
                    next_para.clear()
                    next_para.add_run(
                        'This project addresses the challenge of building a flexible and extensible Multi-Agent System (MAS) '
                        'for automated data processing, feature engineering, machine learning predictions, and visualization. '
                        'The system demonstrates how specialized software agents can collaborate to complete a complete machine '
                        'learning pipeline from raw data ingestion to prediction and visualization.'
                    )
        
        # Update paragraph if changed
        if text != para.text:
            para.clear()
            para.add_run(text)
    
    # Fill tables with data
    for table in doc.tables:
        # Look for dataset information table
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.lower()
                
                # Fill dataset name
                if 'dataset name' in cell_text or 'sample_data' in cell_text:
                    if cell_idx < len(row.cells) - 1:
                        row.cells[cell_idx + 1].text = 'sample_data.csv'
                
                # Fill other dataset attributes
                if 'total rows' in cell_text:
                    if cell_idx < len(row.cells) - 1:
                        row.cells[cell_idx + 1].text = str(results['data_rows'])
                
                if 'total columns' in cell_text:
                    if cell_idx < len(row.cells) - 1:
                        row.cells[cell_idx + 1].text = str(results['data_columns'])
                
                # Replace placeholders
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Save final report
    output_path = "Report_Final.docx"
    doc.save(output_path)
    
    print(f"\n{'='*60}")
    print("Final Report Created Successfully!")
    print(f"{'='*60}")
    print(f"Output file: {output_path}")
    print(f"\nReport includes:")
    print(f"  - Date: {results['date']}")
    print(f"  - Dataset: {results['data_rows']} rows × {results['data_columns']} columns")
    print(f"  - Features: {results['features_count']}")
    print(f"  - Model: {results['model_type']}")
    print(f"  - Predictions: {results['predictions_count']}")
    print(f"{'='*60}")

if __name__ == "__main__":
    create_final_report()

